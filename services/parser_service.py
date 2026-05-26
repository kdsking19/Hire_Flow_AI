import fitz  # PyMuPDF
from docx import Document
import os

def extract_pdf_text(path):
    """Extract all text content from a PDF file."""
    text = ""
    try:
        pdf = fitz.open(path)
        for page in pdf:
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"
        pdf.close()
    except Exception as e:
        print(f"Error parsing PDF at {path}: {e}")
    return text

def extract_docx_text(path):
    """Extract all text content from a DOCX file, including paragraphs and tables."""
    text = ""
    try:
        doc = Document(path)
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text += " | ".join(row_text) + "\n"
    except Exception as e:
        print(f"Error parsing DOCX at {path}: {e}")
    return text

def extract_text(path):
    """Determine the file type and extract its text content."""
    if not os.path.exists(path):
        return ""
    
    _, ext = os.path.splitext(path.lower())
    if ext == ".pdf":
        return extract_pdf_text(path)
    elif ext in [".docx", ".doc"]:
        return extract_docx_text(path)
    else:
        # Attempt fallback to reading as a plain text file
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""
